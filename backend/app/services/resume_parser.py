"""Resume parser — extract structured content from PDF / DOCX / TXT files.

Optimised per resume_parser_optimization.md:
  P0-1: Ordered block traversal (paragraphs & tables in document.xml order)
  P0-2: No hard char truncation — return full text + structured blocks
  P0-3: Structured output: text + blocks + metadata + quality check

Architecture:
  Upload → Format Detector → Document Parser → Structured Blocks
       → Sanity Check → {text, blocks, metadata, quality} → LLM
"""

from __future__ import annotations

import io
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# File magic bytes for format detection
_OLE2_MAGIC = b"\xd0\xcf\x11\xe0"      # Old .doc format (OLE2 compound document)
_ZIP_MAGIC = b"PK\x03\x04"              # .docx format (ZIP/Open XML)

# Sensible char limit (no hard truncation — LLM gets the full text;
# this is a safety cap against malformed files returning millions of chars)
MAX_CHARS = 12000


# ════════════════════════════════════════════════════════════════
# Public API
# ════════════════════════════════════════════════════════════════

async def parse_resume(file_bytes: bytes, filename: str) -> dict:
    """Extract structured content from a resume file.

    Returns:
        dict with keys:
            filename    — original filename
            format      — detected format (pdf/docx/txt)
            text        — plain-text representation (full, up to MAX_CHARS)
            blocks      — ordered list of {type, text|rows, style}
            metadata    — {paragraph_count, table_count, char_count, header_footer_text}
            quality     — {suspicious: bool, issues: [...]}
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        result = _parse_pdf(file_bytes, filename)
    elif ext in (".docx", ".doc"):
        result = _parse_doc_or_docx(file_bytes, filename)
    elif ext == ".txt":
        text = file_bytes.decode("utf-8", errors="replace")
        result = _make_result(filename, "txt", text, [], quality=None)
    else:
        raise ValueError(f"不支持的文件格式: {ext}，请上传 PDF、DOCX 或 TXT 文件。")

    # ── Sanity check ─────────────────────────────────────
    quality = _sanity_check(result["text"], result["blocks"])
    result["quality"] = quality

    # Log quality issues
    if quality["suspicious"]:
        logger.warning(
            "Resume extraction quality suspicious for %r: %s",
            filename, quality["issues"],
        )

    return result


def parse_resume_text(file_bytes: bytes, filename: str) -> str:
    """Convenience wrapper: return plain text only (backward-compatible)."""
    result = _parse_sync(file_bytes, filename)
    return result["text"]


def _parse_sync(file_bytes: bytes, filename: str) -> dict:
    """Synchronous internal entry point (used when no async needed)."""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return _parse_pdf(file_bytes, filename)
    elif ext in (".docx", ".doc"):
        return _parse_doc_or_docx(file_bytes, filename)
    elif ext == ".txt":
        text = file_bytes.decode("utf-8", errors="replace")
        return _make_result(filename, "txt", text, [], quality=None)
    else:
        raise ValueError(f"不支持的文件格式: {ext}，请上传 PDF、DOCX 或 TXT 文件。")


# ════════════════════════════════════════════════════════════════
# PDF
# ════════════════════════════════════════════════════════════════

def _parse_pdf(file_bytes: bytes, filename: str) -> dict:
    """Extract text from PDF bytes."""
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise ImportError("PDF 解析需要 PyPDF2。请运行: pip install PyPDF2")

    reader = PdfReader(io.BytesIO(file_bytes))
    blocks: list[dict] = []
    all_lines: list[str] = []

    for page in reader.pages:
        text = page.extract_text()
        if text:
            all_lines.append(text)
            blocks.append({"type": "paragraph", "text": text, "style": None})

    full_text = "\n".join(all_lines)

    if not full_text.strip():
        raise ValueError("无法从文件中提取文字，文件可能为空白、图片格式或加密文档。")

    return _make_result(
        filename, "pdf",
        text=full_text,
        blocks=blocks,
        char_count=len(full_text),
    )


# ════════════════════════════════════════════════════════════════
# DOC / DOCX
# ════════════════════════════════════════════════════════════════

def _parse_doc_or_docx(file_bytes: bytes, filename: str) -> dict:
    """Extract structured blocks from .doc or .docx in document order."""
    magic = file_bytes[:4]

    # Detect old .doc format (OLE2)
    if magic[:4] == _OLE2_MAGIC:
        raise ValueError(
            "检测到旧版 .doc 格式（Word 2003 及更早版本）。"
            "请用 Microsoft Word 或 WPS 打开文件，"
            "另存为 .docx 格式（Word 2007+），然后重新上传。"
        )

    # Check if it's a valid ZIP/DOCX
    if magic[:4] != _ZIP_MAGIC:
        actual = magic.hex()
        logger.warning(f"Unknown docx magic bytes: {actual} for file {filename!r}")
        raise ValueError(
            f"文件不是有效的 .docx 格式。请确认：\n"
            f"1. 确保文件是用 Word/WPS 保存的标准 .docx 格式\n"
            f"2. 不要直接修改文件后缀名\n"
            f"3. 尝试用 Word 打开后「另存为」→「Word 文档 (*.docx)」"
        )

    try:
        from docx import Document
    except ImportError:
        raise ImportError("DOCX 解析需要 python-docx。请运行: pip install python-docx")

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        logger.warning(f"python-docx failed to open {filename!r}: {e}")
        raise ValueError(
            f"无法打开 .docx 文件：{e}\n\n"
            f"常见原因及解决方法：\n"
            f"1. 文件由 WPS 或其他工具生成，格式不标准 → 用 Word 重新保存\n"
            f"2. 文件带有敏感标签或加密 → 去掉保护后重新保存\n"
            f"3. 文件已损坏 → 尝试重新下载或另存"
        )

    # ── P0-1: Ordered block traversal (doc §2) ──────────
    blocks = _extract_ordered_blocks(doc)

    # ── Extract headers & footers (doc §8) ──────────────
    header_text = _extract_headers_footers(doc)

    # ── Build plain-text representation ─────────────────
    text_lines: list[str] = []
    para_count = 0
    table_count = 0

    for block in blocks:
        if block["type"] == "paragraph":
            text_lines.append(block["text"])
            para_count += 1
        elif block["type"] == "table":
            table_count += 1
            for row in block.get("rows", []):
                text_lines.append(" | ".join(row))
            text_lines.append("")  # blank line after table

    if header_text:
        text_lines.insert(0, f"[Header]\n{header_text}")

    full_text = "\n".join(text_lines)
    char_count = len(full_text)

    if not full_text.strip():
        raise ValueError("无法从文件中提取文字，文件可能为空白、图片格式或加密文档。")

    logger.info(
        "Parsed DOCX %r: %d chars, %d paragraphs, %d tables, %d blocks (ordered)",
        filename, char_count, para_count, table_count, len(blocks),
    )

    return _make_result(
        filename, "docx",
        text=full_text,
        blocks=blocks,
        char_count=char_count,
        paragraph_count=para_count,
        table_count=table_count,
        header_footer_text=header_text if header_text else None,
    )


# ── Ordered Block Extraction (P0-1) ──────────────────────

def _extract_ordered_blocks(doc) -> list[dict]:
    """Yield paragraphs and tables in original document.xml order.

    Per resume_parser_optimization.md §2: this is the single most important fix.
    The old approach (all paragraphs → all tables) destroys the relationship
    between section headings and their tables.
    """
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P

    blocks: list[dict] = []

    # Traverse document body in XML element order
    parent = doc.element.body

    for child in parent.iterchildren():
        if isinstance(child, CT_P):
            para = Paragraph(child, doc)
            text = para.text.strip()
            if not text:
                continue
            blocks.append({
                "type": "paragraph",
                "text": text,
                "style": para.style.name if para.style else None,
            })

        elif isinstance(child, CT_Tbl):
            table = DocxTable(child, doc)
            rows: list[list[str]] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows.append(cells)
            if rows:
                blocks.append({
                    "type": "table",
                    "rows": rows,
                })

    return blocks


# ── Header / Footer Extraction (doc §8) ──────────────────

def _extract_headers_footers(doc) -> str:
    """Extract text from headers and footers across all sections.

    Many resumes put contact info (email, phone, GitHub) in the header.
    """
    lines: list[str] = []

    for section in doc.sections:
        # Header
        try:
            for para in section.header.paragraphs:
                text = para.text.strip()
                if text:
                    lines.append(text)
        except Exception:
            pass  # Some headers may be inaccessible

        # Footer
        try:
            for para in section.footer.paragraphs:
                text = para.text.strip()
                if text:
                    lines.append(text)
        except Exception:
            pass

    return "\n".join(lines)


# ── Sanity Check (doc §9) ────────────────────────────────

def _sanity_check(text: str, blocks: list[dict]) -> dict:
    """Lightweight extraction quality detection.

    Returns {"suspicious": bool, "issues": [...]}
    Per resume_parser_optimization.md §9.
    """
    issues: list[str] = []

    stripped = text.strip()

    # Too little text extracted
    if len(stripped) < 100:
        issues.append("too_little_text")

    # No blocks extracted (for docx)
    if len(blocks) == 0:
        issues.append("no_blocks")

    # Abnormal spacing (suggests broken layout)
    if stripped.count("\n\n\n") > 5:
        issues.append("abnormal_spacing")

    # Very short paragraphs dominating (suggests fragmented extraction)
    if blocks:
        short_paras = sum(
            1 for b in blocks
            if b["type"] == "paragraph" and len(b.get("text", "")) < 10
        )
        if short_paras > len(blocks) * 0.5:  # >50% very short paragraphs
            issues.append("too_many_short_paragraphs")

    return {
        "suspicious": bool(issues),
        "issues": issues,
    }


# ── Result Builder ───────────────────────────────────────

def _make_result(
    filename: str,
    fmt: str,
    text: str,
    blocks: list[dict],
    quality: dict | None = None,
    char_count: int | None = None,
    paragraph_count: int | None = None,
    table_count: int | None = None,
    header_footer_text: str | None = None,
) -> dict:
    """Build the standardised parser output dict.

    Per resume_parser_optimization.md §11:
      - text: plain-text representation (up to MAX_CHARS safety cap)
      - blocks: ordered list of {type, text|rows, style}
      - metadata: extraction stats
      - quality: sanity check result
    """
    if char_count is None:
        char_count = len(text)

    # Count from blocks if not provided
    if paragraph_count is None:
        paragraph_count = sum(1 for b in blocks if b["type"] == "paragraph")
    if table_count is None:
        table_count = sum(1 for b in blocks if b["type"] == "table")

    metadata = {
        "paragraph_count": paragraph_count,
        "table_count": table_count,
        "char_count": char_count,
    }
    if header_footer_text:
        metadata["header_footer_text"] = header_footer_text

    if quality is None:
        quality = {"suspicious": False, "issues": []}

    return {
        "filename": filename,
        "format": fmt,
        "text": text[:MAX_CHARS] if len(text) > MAX_CHARS else text,
        "text_truncated": len(text) > MAX_CHARS,
        "blocks": blocks,
        "metadata": metadata,
        "quality": quality,
    }
